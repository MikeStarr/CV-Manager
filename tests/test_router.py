from unittest.mock import MagicMock, patch

import pytest

from cv_manager.brain import generate_cv_content


def test_generate_cv_content_local_success():
    """Ensure Local LLM routing works and uses the custom client if provided."""
    mock_client = MagicMock()
    mock_response = MagicMock()
    mock_response.choices[0].message.content = "local response"
    mock_client.chat.completions.create.return_value = mock_response

    res = generate_cv_content(system_prompt="sys", user_prompt="user", provider="Local", client=mock_client)
    assert res == "local response"
    mock_client.chat.completions.create.assert_called_once()


def test_generate_cv_content_missing_api_keys(monkeypatch):
    """Ensure ValueError is raised if api keys are missing for cloud providers."""
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    monkeypatch.delenv("XAI_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    with pytest.raises(ValueError, match="DeepSeek API key is missing"):
        generate_cv_content("sys", "user", "DeepSeek")

    with pytest.raises(ValueError, match="Grok API key is missing"):
        generate_cv_content("sys", "user", "Grok")

    with pytest.raises(ValueError, match="OpenAI API key is missing"):
        generate_cv_content("sys", "user", "ChatGPT")


@patch("cv_manager.brain.OpenAI")
def test_generate_cv_content_chatgpt_success(mock_openai_class, monkeypatch):
    """Ensure ChatGPT API is called with correct parameters."""
    monkeypatch.setenv("OPENAI_API_KEY", "mock-openai-key")

    mock_client = MagicMock()
    mock_openai_class.return_value = mock_client

    mock_response = MagicMock()
    mock_response.choices[0].message.content = "chatgpt response"
    mock_client.chat.completions.create.return_value = mock_response

    res = generate_cv_content("sys", "user", "ChatGPT")
    assert res == "chatgpt response"

    mock_openai_class.assert_called_once_with(api_key="mock-openai-key", base_url="https://api.openai.com/v1")
    mock_client.chat.completions.create.assert_called_once_with(
        model="gpt-5.4",
        messages=[{"role": "system", "content": "sys"}, {"role": "user", "content": "user"}],
        temperature=0,
        timeout=15.0,
    )


@patch("cv_manager.brain.OpenAI")
def test_generate_cv_content_deepseek_success(mock_openai_class, monkeypatch):
    """Ensure DeepSeek API is called with correct parameters."""
    monkeypatch.setenv("DEEPSEEK_API_KEY", "mock-ds-key")

    mock_client = MagicMock()
    mock_openai_class.return_value = mock_client

    mock_response = MagicMock()
    mock_response.choices[0].message.content = "deepseek response"
    mock_client.chat.completions.create.return_value = mock_response

    res = generate_cv_content("sys", "user", "DeepSeek")
    assert res == "deepseek response"

    mock_openai_class.assert_called_once_with(api_key="mock-ds-key", base_url="https://api.deepseek.com")
    mock_client.chat.completions.create.assert_called_once_with(
        model="deepseek-chat",
        messages=[{"role": "system", "content": "sys"}, {"role": "user", "content": "user"}],
        temperature=0,
        timeout=15.0,
    )


def test_cv_brain_strips_section_prefixes():
    """Ensure that CVBrain.generate_tailored_content correctly strips [Section] and **[Section]** prefixes."""
    import json

    from cv_manager.brain import CVBrain

    brain = CVBrain()
    brain.client = MagicMock()
    mock_response = MagicMock()
    mock_response.choices[0].message.content = json.dumps(
        [
            {"original_text": "[General] Digital Delivery Manager", "new_text": "[General] Project Manager"},
            {
                "original_text": "**[Professional Experience]** Old Text",
                "new_text": "**[Professional Experience]** New Text",
            },
        ]
    )
    brain.client.chat.completions.create.return_value = mock_response

    cv_structure = [{"text": "Digital Delivery Manager", "style": "Normal"}, {"text": "Old Text", "style": "Normal"}]

    result = brain.generate_tailored_content(job_spec="test", cv_structure=cv_structure, cv_content_md="")

    assert isinstance(result, dict)
    updates = result["updates"]
    assert len(updates) == 2
    assert updates[0]["original_text"] == "Digital Delivery Manager"
    assert updates[0]["new_text"] == "Project Manager"
    assert updates[1]["original_text"] == "Old Text"
    assert updates[1]["new_text"] == "New Text"


def test_cv_brain_flexible_matching():
    """Ensure that CVBrain.generate_tailored_content validates slight space/dash variations and truncated values."""
    import json

    from cv_manager.brain import CVBrain

    brain = CVBrain()
    brain.client = MagicMock()
    mock_response = MagicMock()
    mock_response.choices[0].message.content = json.dumps(
        [
            {
                "original_text": "Project Manager / LSEG – London, UK April 2023– Present",
                "new_text": "Senior Project Manager / LSEG – London, UK April 2023– Present",
            },
            {
                "original_text": "Certified PRINCE2 Practitioner with 15+ years of exp",
                "new_text": "Certified PRINCE2 Practitioner with 15+ years of experience",
            },
        ]
    )
    brain.client.chat.completions.create.return_value = mock_response

    cv_structure = [
        {"text": "Project Manager / LSEG – London, UK\xa0April\xa02023– Present", "style": "Normal"},
        {"text": "Certified PRINCE2 Practitioner with 15+ years of experience.", "style": "Normal"},
    ]

    result = brain.generate_tailored_content(job_spec="test", cv_structure=cv_structure, cv_content_md="")

    assert isinstance(result, dict)
    updates = result["updates"]
    assert len(updates) == 2
    assert "LSEG" in updates[0]["original_text"]
    assert "PRINCE2" in updates[1]["original_text"]


def test_cv_updater_flexible_matching(tmp_path):
    """Ensure that CVUpdater matches and replaces paragraphs with formatting differences or minor truncation."""
    from docx import Document

    from cv_manager.parser.cv_updater import CVUpdater

    # Create dummy docx file
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Project Manager / LSEG – London, UK\xa0April\xa02023– Present")
    doc.add_paragraph("Certified PRINCE2 Practitioner with 15+ years of experience.")
    doc.save(str(doc_path))

    updater = CVUpdater(str(doc_path))

    # 1. Test space/dash variation replacement
    success1 = updater.update_paragraph_text(
        "Project Manager / LSEG – London, UK April 2023– Present", "Senior Project Manager / LSEG"
    )
    assert success1 is True

    # 2. Test truncated text replacement
    success2 = updater.update_paragraph_text(
        "Certified PRINCE2 Practitioner with 15+ years of exp",
        "Certified PRINCE2 Practitioner with 15+ years of senior experience",
    )
    assert success2 is True

    # Save and read back to verify
    out_path = tmp_path / "out.docx"
    updater.save(str(out_path))

    doc2 = Document(str(out_path))
    paras = [p.text for p in doc2.paragraphs]
    assert "Senior Project Manager / LSEG" in paras
    assert "Certified PRINCE2 Practitioner with 15+ years of senior experience" in paras


def test_cv_updater_partial_substring_replacement(tmp_path):
    """Ensure that CVUpdater replaces ONLY the matching sentence inside a multi-sentence paragraph."""
    from docx import Document

    from cv_manager.parser.cv_updater import CVUpdater

    doc_path = tmp_path / "test_partial.docx"
    doc = Document()
    doc.add_paragraph("Certified PRINCE2 Practitioner. I have 15+ years of experience. Always focusing on delivery.")
    doc.save(str(doc_path))

    updater = CVUpdater(str(doc_path))

    # Replace ONLY the middle sentence
    success = updater.update_paragraph_text(
        "I have 15+ years of experience", "I possess over 15 years of digital delivery experience"
    )
    assert success is True

    out_path = tmp_path / "out_partial.docx"
    updater.save(str(out_path))

    doc2 = Document(str(out_path))
    paras = [p.text for p in doc2.paragraphs]
    # The first and last sentences should be perfectly preserved!
    assert (
        paras[0]
        == "Certified PRINCE2 Practitioner. I possess over 15 years of digital delivery experience. Always focusing on delivery."
    )


def test_cv_updater_preserves_bolding_across_runs(tmp_path):
    """Ensure that CVUpdater preserves bold runs when modifying adjacent text in the same paragraph."""
    from docx import Document

    from cv_manager.parser.cv_updater import CVUpdater

    doc_path = tmp_path / "bold_test.docx"
    doc = Document()
    p = doc.add_paragraph()
    # Add a bold run and a normal run
    r0 = p.add_run("GTM Readiness:")
    r0.bold = True
    r1 = p.add_run(" Led cross-functional team of 10 people.")
    r1.bold = False
    doc.save(str(doc_path))

    updater = CVUpdater(str(doc_path))
    success = updater.update_paragraph_text(
        "Led cross-functional team of 10 people", "Led cross-functional team of 12 senior engineers"
    )
    assert success is True

    out_path = tmp_path / "bold_out.docx"
    updater.save(str(out_path))

    doc2 = Document(str(out_path))
    p2 = doc2.paragraphs[0]

    # Assert formatting is preserved
    assert len(p2.runs) >= 2
    # The first run should still be bold and contain "GTM Readiness"
    assert p2.runs[0].bold is True
    assert "GTM Readiness" in p2.runs[0].text
    # The updated text should contain the new phrasing
    assert "12 senior engineers" in p2.text
