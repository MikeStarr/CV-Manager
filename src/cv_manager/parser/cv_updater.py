import os
import re
import difflib
from docx import Document

class CVUpdater:
    """Updates a .docx file while preserving paragraph styles."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)

    def replace_para_text_preserving_runs(self, para, new_p_text: str):
        """
        Replaces the entire paragraph text with new_p_text,
        aligning the old and new text to preserve run-level formatting.
        """
        runs = para.runs
        if not runs:
            para.text = new_p_text
            return

        old_p_text = para.text

        # Calculate run boundaries in original text
        run_boundaries = []
        curr = 0
        for r in runs:
            r_len = len(r.text)
            run_boundaries.append((curr, curr + r_len))
            curr += r_len

        # Map each run index to its character range in old_p_text
        # Initialize new text list for each run
        new_run_texts = {i: [] for i in range(len(runs))}

        matcher = difflib.SequenceMatcher(None, old_p_text, new_p_text)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                # Map characters directly
                for x in range(i1, i2):
                    # Find which run contains index x
                    for idx, (r_start, r_end) in enumerate(run_boundaries):
                        if r_start <= x < r_end:
                            new_run_texts[idx].append(new_p_text[j1 + (x - i1)])
                            break
            elif tag == 'replace':
                N = i2 - i1
                M = j2 - j1
                if M > 0 and N > 0:
                    for k in range(M):
                        # Distribute character k of replacement
                        orig_idx = i1 + int((k / M) * N)
                        if orig_idx >= i2:
                            orig_idx = i2 - 1
                        # Find which run contains orig_idx
                        for idx, (r_start, r_end) in enumerate(run_boundaries):
                            if r_start <= orig_idx < r_end:
                                new_run_texts[idx].append(new_p_text[j1 + k])
                                break
                elif M > 0:
                    # Treat as insert
                    # Assign to the run containing i1 (or the last run if at the end)
                    assigned = False
                    for idx, (r_start, r_end) in enumerate(run_boundaries):
                        if r_start <= i1 < r_end:
                            new_run_texts[idx].append(new_p_text[j1:j2])
                            assigned = True
                            break
                    if not assigned:
                        new_run_texts[len(runs) - 1].append(new_p_text[j1:j2])
            elif tag == 'insert':
                # Find which run contains i1
                assigned = False
                for idx, (r_start, r_end) in enumerate(run_boundaries):
                    if r_start <= i1 < r_end:
                        new_run_texts[idx].append(new_p_text[j1:j2])
                        assigned = True
                        break
                if not assigned:
                    new_run_texts[len(runs) - 1].append(new_p_text[j1:j2])
            # 'delete' does not add any text

        # Apply the new text to each run
        for idx, r in enumerate(runs):
            r.text = "".join(new_run_texts[idx])

    def update_paragraph_text(self, target_text: str, new_text: str) -> bool:
        """
        Finds a paragraph containing target_text (with flexible spacing/dashes)
        and replaces the matched portion with new_text while preserving formatting.
        """
        target_clean = target_text.strip()
        if not target_clean:
            return False

        # Formatting helper
        def format_new_text(val: str) -> str:
            if ":" in val:
                prefix, rest = val.split(":", 1)
                if prefix.lower().strip() == "areas of expertise":
                    if not prefix.startswith("*"):
                        return f"*{prefix}*: {rest.strip()}"
            return val

        new_text_fmt = format_new_text(new_text)

        # Build regex for flexible whitespace and dashes
        escaped = re.escape(target_clean)
        pattern = re.sub(r"\\s+", lambda m: r"[\s\xa0]+", escaped)
        pattern = re.sub(r"\\-|\\–|\\—", lambda m: r"[-–—]", pattern)
        try:
            regex = re.compile(pattern, re.IGNORECASE)
        except Exception:
            regex = None

        # Step 1: Regex Paragraph Match (includes run-level and multi-run matches)
        if regex:
            for para in self.doc.paragraphs:
                match = regex.search(para.text)
                if match:
                    start, end = match.span()
                    match_len = end - start
                    # If the match covers 80% or more of the paragraph, replace the whole paragraph
                    if match_len >= 0.80 * len(para.text):
                        self.replace_para_text_preserving_runs(para, new_text_fmt)
                    else:
                        new_p_text = para.text[:start] + new_text_fmt + para.text[end:]
                        self.replace_para_text_preserving_runs(para, new_p_text)
                    return True

        # Step 2: Exact Match of entire paragraph text
        for para in self.doc.paragraphs:
            if para.text.strip() == target_clean:
                self.replace_para_text_preserving_runs(para, new_text_fmt)
                return True

        # Step 3: Normalized Substring Match fallback
        def normalize(t: str) -> str:
            t = re.sub(r"\s+", " ", t)
            t = re.sub(r"[-–—]", "-", t)
            return t.strip().strip(".,;:?!-–—•*\"'").lower()

        target_norm = normalize(target_clean)
        if len(target_norm) >= 15:
            for para in self.doc.paragraphs:
                p_norm = normalize(para.text)
                idx = p_norm.find(target_norm)
                if idx != -1:
                    # We need to map target_norm back to the original index in para.text
                    matcher = difflib.SequenceMatcher(None, p_norm, para.text.lower())
                    start_char = -1
                    end_char = -1
                    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                        if tag == 'equal':
                            if i1 <= idx < i2:
                                start_char = j1 + (idx - i1)
                            if i1 < idx + len(target_norm) <= i2:
                                end_char = j1 + (idx + len(target_norm) - i1)

                    if start_char == -1 or end_char == -1:
                        if len(target_norm) >= 0.80 * len(p_norm):
                            self.replace_para_text_preserving_runs(para, new_text_fmt)
                        else:
                            words_count = len(target_clean.split())
                            para_words = para.text.split()
                            if p_norm.startswith(target_norm):
                                boundary_char = len(" ".join(para_words[:words_count]))
                                self.replace_para_text_preserving_runs(para, para.text[:boundary_char].replace(para.text[:boundary_char], new_text_fmt) + para.text[boundary_char:])
                            elif p_norm.endswith(target_norm):
                                boundary_char = len(para.text) - len(" ".join(para_words[-words_count:]))
                                self.replace_para_text_preserving_runs(para, para.text[:boundary_char] + new_text_fmt)
                            else:
                                continue
                    else:
                        match_len = end_char - start_char
                        if match_len >= 0.80 * len(para.text):
                            self.replace_para_text_preserving_runs(para, new_text_fmt)
                        else:
                            new_p_text = para.text[:start_char] + new_text_fmt + para.text[end_char:]
                            self.replace_para_text_preserving_runs(para, new_p_text)
                    return True

            # Step 4: Fuzzy Match on entire paragraph
            best_para = None
            best_ratio = 0.0
            for para in self.doc.paragraphs:
                p_norm = normalize(para.text)
                ratio = difflib.SequenceMatcher(None, target_norm, p_norm).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_para = para
            
            if best_ratio >= 0.80 and best_para is not None:
                self.replace_para_text_preserving_runs(best_para, new_text_fmt)
                return True

    def save(self, output_path: str = None):
        """Saves the modified document."""
        target = output_path if output_path else self.file_path
        self.doc.save(target)

if __name__ == "__main__":
    # Test the updater (requires a sample docx in cvs/)
    import os
    test_file = "cvs/test_cv.docx"
    if os.path.exists(test_file):
        updater = CVUpdater(test_file)
        # This is a very simple test: try to replace a known string
        # In a real scenario, we'd use the parsed structure from Phase 1
        success = updater.update_paragraph_text("Old Text", "New Updated Text")
        if success:
            updater.save(test_file)
            print("Update successful")
        else:
            print("Target text not found")
    else:
        print("No test file found in cv.test_cv.docx")
