"""Utility for loading Word CV files.

The function :func:`load_cv_text` uses the ``python-docx`` package to read a
.docx file and return its plain text.  The implementation is defensive – if
the library is missing or the file cannot be parsed, an exception is raised
so that the caller can handle it gracefully.
"""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
except Exception as exc:  # pragma: no cover - defensive
    raise RuntimeError("python-docx is required to read CV files") from exc


def load_cv_text(file_path: str | Path) -> str:
    """Return the concatenated text of all paragraphs in a .docx file.

    Parameters
    ----------
    file_path:
        Path to the Word document.
    """
    doc = Document(str(file_path))
    return "\n".join(p.text for p in doc.paragraphs)


from typing import Any

from docx import Document


class CVParser:
    """Parses a .docx CV to extract its structural elements."""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)

    def parse(self) -> list[dict[str, Any]]:
        """
        Parses the document and returns a list of paragraphs with their styles.
        Each dict contains:
        - text: The text content of the paragraph
        - style: The style name applied to the paragraph
        - is_heading: Boolean indicating if it's a heading
        """
        elements = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Simple heuristic for headings: check if style name contains 'Heading'
            if para.style is not None:
                is_heading = "Heading" in para.style.name
                style_name = para.style.name
            else:
                is_heading = False
                style_name = "Normal"

            elements.append({"text": text, "style": style_name, "is_heading": is_heading})
        return elements


if __name__ == "__main__":
    # Test the parser (requires a sample docx in cvs/)
    import os

    test_file = "cvs/test_cv.docx"
    if os.path.exists(test_file):
        parser = CVParser(test_file)
        for el in parser.parse():
            print(f"[{el['style']}] {el['text']}")
    else:
        print("No test file found in cvs/test_cv.docx")
